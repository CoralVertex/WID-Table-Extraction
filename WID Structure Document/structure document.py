from bs4 import BeautifulSoup 
from urllib.request import Request,  urlopen
from numpy import NAN, NaN
import pandas as pd
import regex as re
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
from docx.shared import Pt


#function to write multiple dataframes in one sheet in Excel
def multiple_dfs(df_list, sheets, file_name, spaces):
    book = load_workbook(file_name)
    writer = pd.ExcelWriter(file_name,engine='openpyxl')  
    writer.book = book
    writer.sheets = dict((ws.title, ws) for ws in book.worksheets) 
    row = 0
    for dataframe in df_list:
        dataframe.to_excel(writer,sheets,startrow=row , startcol=0, index=False)          
        row = row + len(dataframe.index) + spaces + 1
    writer.save()

def export_to_docx(df_list, title):
    doctitle = title + ".docx"
    document = Document()

    document.add_heading(title, 1)

    sections = document.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)     

    
    
    row = 0
    pagerow = 0
    for dataframe in df_list:      
        #row = len(dataframe.index) 
        row = 1

        if dataframe.columns.values[0] != 'Constraints' and len(dataframe)==1 and pagerow>40:
            document.add_page_break()

            h = document.add_heading(dataframe.columns[0][:-2], 2)
            r = h.add_run()
            if dataframe.columns[0][-2:]=='10' or dataframe.columns[0][-2:]=='11':
                r.add_text(' ')
                r.add_picture('core.png', width=Inches(0.15))
            if dataframe.columns[0][-2:]=='01' or dataframe.columns[0][-2:]=='11':
                r.add_text(' ')
                r.add_picture('trigger.png', width=Inches(0.15))
            table = document.add_table(rows=row, cols=1, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = dataframe.iloc[0]
            hdr_cells[0].width = Inches(6.75)

            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(12)

            pagerow = 1

        elif dataframe.columns.values[0] != 'Constraints' and len(dataframe)==1:
            #p = document.add_paragraph()
            #run = p.add_run()
            #run.add_break()  # by default adds WD_BREAK.LINE
            h = document.add_heading(dataframe.columns[0][:-2], 2)
            r = h.add_run()
            if dataframe.columns[0][-2:]=='10' or dataframe.columns[0][-2:]=='11':
                r.add_text(' ')
                r.add_picture('core.png', width=Inches(0.15))
            if dataframe.columns[0][-2:]=='01' or dataframe.columns[0][-2:]=='11':
                r.add_text(' ')
                r.add_picture('trigger.png', width=Inches(0.15))
            #document.add_heading(dataframe.columns[0][:-2], 1) 
            #if dataframe.columns[0][-2:]=='10' or dataframe.columns[0][-2:]=='11':
            #    document.add_picture('core.png', width=Inches(0.15))
            #if dataframe.columns[0][-2:]=='01' or dataframe.columns[0][-2:]=='11':
            #    document.add_picture('trigger.png', width=Inches(0.15))
            table = document.add_table(rows=row, cols=1, style='Table Grid')
            hdr_cells = table.rows[0].cells
            #hdr_cells[0].text = dataframe.columns[0]
            #row_cells = table.add_row().cells
            hdr_cells[0].text = dataframe.iloc[0]
            #hdr_cells[1].text = dataframe.iloc[0]#['']
            hdr_cells[0].width = Inches(6.75)

            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(12)
            
            pagerow = pagerow + len(table.rows)

        elif dataframe.columns.values[0] == 'FieldName':
            #row = row + 1
            table = document.add_table(rows=row, cols=4, style='Light List Accent 1')
            
            x = 0
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'FieldName'
            hdr_cells[1].text = 'FieldType'
            hdr_cells[2].text = 'Constraint'
            hdr_cells[3].text = 'FieldDesc'
            while x < len(dataframe.index): 
                row_cells = table.add_row().cells
                row_cells[0].text = dataframe.iloc[x]['FieldName']
                row_cells[0].width = Inches(1.5)
                row_cells[1].text = dataframe.iloc[x]['FieldType']
                row_cells[1].width = Inches(1.25)
                row_cells[2].text = dataframe.iloc[x]['Constraint']
                row_cells[2].width = Inches(1.25)
                if pd.isna(dataframe.iloc[x]['FieldDesc']):
                    row_cells[3].text = ''
                    row_cells[3].width = Inches(3)
                else:
                    row_cells[3].text = dataframe.iloc[x]['FieldDesc']
                    row_cells[3].width = Inches(3)
                x = x + 1  
             
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)
            pagerow = pagerow + len(table.rows)

        elif dataframe.columns.values[0] == 'Constraints':
            #row = row + 1
            table = document.add_table(rows=row, cols=1, style='Medium Shading 1 Accent 1')

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Constraints'
            row_cells = table.add_row().cells
            row_cells[0].text = dataframe.iloc[0]['Constraints']
            row_cells[0].width = Inches(6.75)

            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(11)
                        font.bold = False
            pagerow = pagerow + len(table.rows)
            #print(pagerow)
    #document.add_page_break()

    document.save(doctitle)


fieldDetails = pd.read_excel('WID_Fields.xlsx')
tableDescriptions = pd.read_excel('WID_Tables.xlsx')

# match to sql doc

sql_file = 'wid30final.sql'
sql = open(sql_file, 'r', encoding="utf-16")
fullText = sql.read()
allTables = fullText.split('CREATE TABLE [dbo].')
sqlSection = pd.DataFrame(columns=['chunk'])

for x in allTables:
    chunk = pd.DataFrame([{'chunk':x}])
    sqlSection = pd.concat([chunk,sqlSection], ignore_index=True)
    

tableBase = pd.DataFrame(columns=['TableName','FieldName'])
#tableSection = pd.DataFrame(columns=['TableName','OrigFieldName','OrigKeys'])
tableStructure = pd.DataFrame(columns=['TableName','OrigFieldName','OrigKeys'])
tablePrimaryKeys = pd.DataFrame(columns=['TableName','FieldName','Constraint'])
tableForeignKeys = pd.DataFrame(columns=['TableName','BaseFields','ReferenceTable','ReferenceFields','KeyName', 'ConstraintText','ConstraintNum'])
tableForeignKeysCombined = pd.DataFrame(columns=['TableName','ConstraintText'])

sqlSection.reset_index(drop=True)
n = 0
# extract tables
while n < len(sqlSection.index)-1: 
    section = pd.DataFrame.from_dict([{'TableName':sqlSection.chunk.loc[n].split("(\n\t")[0], 'FieldName':sqlSection.chunk.loc[n].split("(\n\t")[1], 'OrigKeys':sqlSection.chunk.loc[n].split("PRIMARY KEY CLUSTERED")[1]}])
    tableBase = pd.concat([section,tableBase], ignore_index=True)
    n=n+1

m = 0
# look for fields within tables
while m < len(tableBase.index):
    splitString = tableBase.FieldName.loc[m].split(",\n")
    tabName = tableBase.TableName.loc[m].replace('[','').replace(']','')
    # clean up keys
    tableBase.OrigKeys.loc[m] = tableBase.OrigKeys.loc[m].split('\n)WITH')[0].replace(' \n(\n\t','')
    # create secondary list of primary keys for join
    constraintString = tableBase.OrigKeys.loc[m].split(",\n")
    fieldName = 0
    for y in constraintString:
        y = y.replace('\t','').replace('[','').replace(']','').replace(', ','').replace('ASC','').replace(' ','')
        pk = pd.DataFrame.from_dict([{'TableName':tabName, 'FieldName':y, 'Constraint':'Primary Key'}])
        tablePrimaryKeys = pd.concat([pk,tablePrimaryKeys], ignore_index=True)

    # split out fields
    rowNum = 0
    for x in splitString:
        if x[:1]=='[' or x[:1]=='\t':
            rowNum = rowNum+1
            x = x.replace('\t','').replace('[','').replace(']','').replace(', ',',')
            ts = pd.DataFrame.from_dict([{'TableName':tabName, 'OrigFieldName':x, 'OrigKeys':NAN, 'FieldName':x.split(' ')[0], 'FieldType':x.split(' ')[1], 'FieldNULL':x.split(' ')[2], 'ColumnNum':str(rowNum)}])
            tableStructure = pd.concat([ts,tableStructure], ignore_index=True)
    m = m + 1

# look for foreign key details at end of document
allForeign = sqlSection.loc[0]['chunk'].split('GO\nALTER TABLE ')
for f in allForeign:
    if re.search('WITH CHECK ADD  CONSTRAINT', f):
        fk = pd.DataFrame.from_dict([{'TableName':f.split('  ')[0].replace('[dbo].[','').replace(']',''),'BaseFields':f.split('FOREIGN KEY(')[1].split(')\n')[0],'ReferenceTable':f.split('REFERENCES ')[1].split('(')[0],'ReferenceFields':f.split('REFERENCES ')[1].split('(')[1].replace(')\n',''),'KeyName':f.split('WITH CHECK ADD  CONSTRAINT [')[1].split('] FOREIGN KEY')[0]}])
        tableForeignKeys = pd.concat([fk,tableForeignKeys], ignore_index=True)
# generate foreign key text for document
l = 0
tname2 = ''    
r = 1
while l < len(tableForeignKeys.index):
    tname = tableForeignKeys.loc[l]['TableName']
    if tname == tname2:
        r = r + 1
    else:
        r = 1
    tableForeignKeys.loc[l, 'ConstraintText'] = str(r) + '. Foreign Key (' + tableForeignKeys.loc[l]['BaseFields'].replace('[',tname+'.').replace(']','') + ') references (' + tableForeignKeys.loc[l]['ReferenceFields'].replace(']','').replace('[',tableForeignKeys.loc[l]['ReferenceTable'].replace('] ','').replace('[dbo].[','')+'.') + ')'
    if tname == tname2:
        cname = tableForeignKeys.loc[l]['ConstraintText']
        newname = tableForeignKeysCombined[tableForeignKeysCombined['TableName']==tname]['ConstraintText'].values[0] + '\n' + cname
        tableForeignKeysCombined.loc[tableForeignKeysCombined['TableName']==tname, 'ConstraintText'] = newname
    else:
        fkc = pd.DataFrame([{'TableName':tname, 'ConstraintText':tableForeignKeys.loc[l]['ConstraintText']}])
        tableForeignKeysCombined = pd.concat([fkc,tableForeignKeysCombined], ignore_index=True)
    tname2 = tableForeignKeys.loc[l]['TableName']
    l = l + 1


#display(tablePrimaryKeys)
#display(tableForeignKeysCombined)

##add end row
ts = pd.DataFrame([{'TableName':"end",'orig_TableName':"end", 'OrigFieldName':"end", 'OrigKeys':NAN, 'FieldName':"end", 'FieldType':NAN, 'FieldNULL':NAN, 'ColumnNum':'1'}])
tableStructure = pd.concat([ts,tableStructure], ignore_index=True)


# join to field details
tableStructure['orig_TableName'] = tableStructure['TableName']
tableStructure['orig_FieldName'] = tableStructure['FieldName']
tableStructure['TableName'] = tableStructure['TableName'].str.lower()
tableStructure['FieldName'] = tableStructure['FieldName'].str.lower()
tableStructure = pd.merge(tableStructure,fieldDetails, left_on=['TableName','FieldName'], right_on=['TableName','FieldName'], suffixes = ('_structure','_detail'), how = 'left')

## join to table Primary Keys to populate initial "Constraint" in table
tablePrimaryKeys['TableName'] = tablePrimaryKeys['TableName'].str.lower()
tablePrimaryKeys['FieldName'] = tablePrimaryKeys['FieldName'].str.lower()

tableStructure = pd.merge(tableStructure,tablePrimaryKeys, how='outer', left_on=['TableName','FieldName'], right_on=['TableName','FieldName'])
tableStructure = tableStructure.sort_values(by=['TableName','FieldOrder'], ascending=[True,False])

## look for field name in tableForeignKeys to get number of foreign key for table
f = 0

while f < len(tableStructure.index):
    tname = tableStructure.loc[f]['orig_TableName']
    fname = tableStructure.loc[f]['orig_FieldName']
    tempDf = tableForeignKeys[tableForeignKeys['TableName']==tname].reset_index()
    if len(tempDf.index)>0:
        for i in tempDf.index: 
            tempDf.at[i,'ConstraintNum']=tempDf.loc[i]['ConstraintText'][:1]
    keystring = ''
    k = 0
    while k < len(tempDf.index)-1:
        if bool(re.search(fname, tempDf.loc[k]['BaseFields']))==True:
            if keystring == '':
                keystring = tempDf.loc[k]['ConstraintNum']
            else:
                keystring = keystring + ',' + tempDf.loc[k]['ConstraintNum']
        k = k + 1
    if pd.isna(tableStructure.loc[f]['Constraint']): 
        tableStructure.at[f,'Constraint'] = keystring
    else:
        tableStructure.at[f,'Constraint'] = tableStructure.loc[f]['Constraint'] +'\n' + keystring
    f = f + 1
    k = 0

## join to tableForeignKeysCombined to get foreign key full text

formattedTable = pd.DataFrame(columns=['TableName','FieldName','FieldType','Constraint','FieldDesc'])
#dflist = pd.DataFrame(columns=['tableType','tableHeader','tableDetails','tableFooter'])
z = 0
tname2 = tableStructure.loc[0]['TableName']    

#try to fix order problem
tableStructure = tableStructure.sort_values(by=['TableName','FieldOrder'], ascending=[True,False]).reset_index()

#merge and drop duplicates - start here
dflist = pd.DataFrame(columns=['tabletype','tablename','sequence','content'])

while z < len(tableStructure.index):
    tname = tableStructure.loc[z]['TableName']
    if len(tableDescriptions[tableDescriptions['TableName']==tname].reset_index().index)>0:
        tdesc = tableDescriptions[tableDescriptions['TableName']==tname].reset_index().loc[0]['TableDesc']
        ttype = tableDescriptions[tableDescriptions['TableName']==tname].reset_index().loc[0].at['TableType']
        ctable = str(int(tableDescriptions[tableDescriptions['TableName']==tname].reset_index().loc[0].at['CoreTable']))
        ctrigger = str(int(tableDescriptions[tableDescriptions['TableName']==tname].reset_index().loc[0].at['Trigger']))
    else:
        tdesc = ''
        ttype = ''
    theader = pd.DataFrame([tdesc], columns=[tableStructure.loc[z]['orig_TableName']+ctable+ctrigger])
    
    fieldname = tableStructure.loc[z]['ColumnNum'] + '. ' + tableStructure.loc[z]['orig_FieldName']
    ft = pd.DataFrame([{'TableName':tname, 'FieldName':fieldname, 'FieldType':tableStructure.loc[z]['FieldType'], 'Constraint':tableStructure.loc[z]['Constraint'], 'FieldDesc':tableStructure.loc[z]['FieldDesc']}])
    formattedTable = pd.concat([ft,formattedTable], ignore_index=True)

    if tname != tname2:

        if len(tableDescriptions[tableDescriptions['TableName']==tname2].reset_index().index)>0:
            ttype = tableDescriptions[tableDescriptions['TableName']==tname2].reset_index().loc[0].at['TableType']
        else:
            ttype = ''

        tt = pd.DataFrame([{'tabletype':ttype,'tablename':tname2,'sequence':1,'content':theader2}])    
        dflist = pd.concat([tt,dflist], ignore_index=True)
        tt2 = pd.DataFrame([{'tabletype':ttype,'tablename':tname2,'sequence':2,'content':formattedTable[formattedTable['TableName']==tname2][['FieldName','FieldType','Constraint','FieldDesc']]}])
        dflist = pd.concat([tt2, dflist], ignore_index=True)
        if len(tableForeignKeysCombined[tableForeignKeysCombined['TableName'].str.lower()==tname2].reset_index().index)>0:
            fkey = pd.DataFrame([tableForeignKeysCombined[tableForeignKeysCombined['TableName'].str.lower()==tname2].reset_index().loc[0]['ConstraintText']], columns=['Constraints'])
            dfl = pd.DataFrame([{'tabletype':ttype,'tablename':tname2,'sequence':3,'content':fkey}])
            dflist = pd.concat([dfl, dflist], ignore_index=True)


    tname2 = tableStructure.loc[z]['TableName']
    theader2 = pd.DataFrame([tdesc], columns=[tableStructure.loc[z]['orig_TableName']+ctable+ctrigger])
    #theader2 = pd.DataFrame(["<w:style = 'Heading 2'>" + tableStructure.loc[z]['orig_TableName'] + '</w:style>' , tdesc], columns=['heading','desc'])
    z = z + 1

tt2 = pd.DataFrame([{'tabletype':ttype,'tablename':tname2,'sequence':2,'content':formattedTable[formattedTable['TableName']==tname2][['FieldName','FieldType','Constraint','FieldDesc']]}])
dflist = pd.concat([tt2, dflist], ignore_index=True)
if len(tableForeignKeysCombined[tableForeignKeysCombined['TableName'].str.lower()==tname2].reset_index().index)>0:
    fkey = pd.DataFrame([tableForeignKeysCombined[tableForeignKeysCombined['TableName'].str.lower()==tname2].reset_index().loc[0]['ConstraintText']], columns=['Constraints'])
    dfl = pd.DataFrame([{'tabletype':ttype,'tablename':tname2,'sequence':3,'content':fkey}])
    dflist = pd.concat([dfl, dflist], ignore_index=True)


display(dflist[dflist['tabletype']=='lookup'].sort_values(by=['tablename','sequence']))

#multiple_dfs(dflist[dflist['tabletype']=='lookup'].sort_values(by=['tablename','sequence'])['content'], 'lookup', 'structure.xlsx', 0)
#multiple_dfs(dflist[dflist['tabletype']=='data'].sort_values(by=['tablename','sequence'])['content'], 'data', 'structure.xlsx', 0)
#multiple_dfs(dflist[dflist['tabletype']=='xwalk'].sort_values(by=['tablename','sequence'])['content'], 'xwalk', 'structure.xlsx', 0)
#multiple_dfs(dflist[dflist['tabletype']=='admin'].sort_values(by=['tablename','sequence'])['content'], 'admin', 'structure.xlsx', 0)


export_to_docx(dflist[dflist['tabletype']=='lookup'].sort_values(by=['tablename','sequence'])['content'], 'Look-Up Tables')
export_to_docx(dflist[dflist['tabletype']=='data'].sort_values(by=['tablename','sequence'])['content'], 'Data Tables')
export_to_docx(dflist[dflist['tabletype']=='xwalk'].sort_values(by=['tablename','sequence'])['content'], 'Crosswalk Tables')
export_to_docx(dflist[dflist['tabletype']=='admin'].sort_values(by=['tablename','sequence'])['content'], 'Administrative Tables')
print('Finished!')
